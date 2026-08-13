from django.http import JsonResponse, HttpResponse
from django.views.decorators.csrf import csrf_exempt

import json
import re
import io

from django.core.mail import send_mail

from rest_framework.decorators import api_view
from rest_framework.response import Response
from rest_framework import status

from .otp import generate_otp, store_otp, verify_otp

from docx import Document
from pypdf import PdfReader

from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4

from .pipeline import simplify_pipeline
from .grammar import grammar_score


# =========================================================
# READABILITY CALCULATION
# =========================================================

def calculate_readability(text):
    """
    Calculate:

    - Flesch Reading Ease
    - Flesch-Kincaid Grade Level

    Higher Reading Ease = easier to read.
    Lower Grade Level = easier to read.
    """

    if not isinstance(text, str):
        text = str(text)

    text = text.strip()

    if not text:
        return {
            "reading_ease": 0,
            "grade_level": 0
        }

    # -----------------------------------------------------
    # Sentence detection
    # -----------------------------------------------------

    sentences = re.split(
        r"[.!?]+",
        text
    )

    sentences = [
        sentence.strip()
        for sentence in sentences
        if sentence.strip()
    ]

    # -----------------------------------------------------
    # Word detection
    # -----------------------------------------------------

    words = re.findall(
        r"\b[a-zA-Z]+(?:'[a-zA-Z]+)?\b",
        text
    )

    if not words:
        return {
            "reading_ease": 0,
            "grade_level": 0
        }

    sentence_count = max(
        len(sentences),
        1
    )

    word_count = len(words)

    # -----------------------------------------------------
    # Syllable counter
    # -----------------------------------------------------

    def count_syllables(word):

        word = word.lower()

        word = re.sub(
            r"[^a-z]",
            "",
            word
        )

        if not word:
            return 1

        vowels = "aeiouy"

        syllables = 0
        previous_vowel = False

        for character in word:

            current_vowel = (
                character in vowels
            )

            if current_vowel and not previous_vowel:
                syllables += 1

            previous_vowel = current_vowel

        # Silent e
        if (
            word.endswith("e")
            and not word.endswith("le")
            and syllables > 1
        ):
            syllables -= 1

        # Words ending in -ed
        if (
            word.endswith("ed")
            and len(word) > 3
            and syllables > 1
        ):

            if (
                len(word) >= 3
                and word[-3] not in vowels
            ):
                syllables -= 1

        return max(
            syllables,
            1
        )

    # -----------------------------------------------------
    # Total syllables
    # -----------------------------------------------------

    total_syllables = sum(
        count_syllables(word)
        for word in words
    )

    # -----------------------------------------------------
    # Average values
    # -----------------------------------------------------

    words_per_sentence = (
        word_count / sentence_count
    )

    syllables_per_word = (
        total_syllables / word_count
    )

    # -----------------------------------------------------
    # Flesch Reading Ease
    # -----------------------------------------------------

    reading_ease = (
        206.835
        - (1.015 * words_per_sentence)
        - (84.6 * syllables_per_word)
    )

    reading_ease = max(
        min(reading_ease, 100),
        -100
    )

    # -----------------------------------------------------
    # Flesch-Kincaid Grade
    # -----------------------------------------------------

    grade_level = (
        (0.39 * words_per_sentence)
        + (11.8 * syllables_per_word)
        - 15.59
    )

    grade_level = max(
        round(grade_level, 1),
        0
    )

    return {
        "reading_ease": round(
            reading_ease,
            2
        ),

        "grade_level": grade_level
    }


# =========================================================
# NORMALIZE DIFFICULT WORDS
# =========================================================

def normalize_difficult_words(words):
    """
    Normalize difficult-word results while PRESERVING:

    - word
    - meaning
    - replacement

    Supported input:

    [
        "ubiquitous"
    ]

    OR:

    [
        {
            "word": "ubiquitous",
            "meaning": "Found everywhere; very common.",
            "replacement": "common"
        }
    ]

    The frontend receives objects so it can display the meaning.
    """

    if not isinstance(words, list):
        return []

    cleaned = []
    seen = set()

    for item in words:

        if isinstance(item, dict):
            word = str(item.get("word", "")).strip()
            meaning = str(
                item.get("meaning", "Meaning not available.")
            ).strip()
            replacement = str(
                item.get("replacement", "")
            ).strip()

        else:
            word = str(item).strip()
            meaning = "Meaning not available."
            replacement = ""

        if not word:
            continue

        key = word.lower()

        if key in seen:
            continue

        seen.add(key)

        if not meaning:
            meaning = "Meaning not available."

        cleaned.append(
            {
                "word": word,
                "meaning": meaning,
                "replacement": replacement,
            }
        )

    return cleaned


# =========================================================
# CLEAN AI WORD MAPPINGS
# =========================================================

def clean_word_mappings(
    mappings,
    original_text,
    simplified_text
):
    """
    Clean and validate vocabulary replacement mappings.

    Expected format:

    [
        {
            "word": "rapid",
            "replacement": "fast"
        }
    ]
    """

    if not isinstance(
        mappings,
        list
    ):
        return []

    cleaned = []

    original_lower = (
        original_text.lower()
    )

    simplified_lower = (
        simplified_text.lower()
    )

    for item in mappings:

        if not isinstance(
            item,
            dict
        ):
            continue

        word = item.get(
            "word",
            ""
        )

        replacement = item.get(
            "replacement",
            ""
        )

        if (
            word is None
            or replacement is None
        ):
            continue

        word = str(
            word
        ).strip()

        replacement = str(
            replacement
        ).strip()

        if (
            not word
            or not replacement
        ):
            continue

        # -------------------------------------------------
        # Ignore identical replacements
        # -------------------------------------------------

        if (
            word.lower()
            == replacement.lower()
        ):
            continue

        # -------------------------------------------------
        # Make sure original term exists
        # -------------------------------------------------

        if (
            word.lower()
            not in original_lower
        ):
            continue

        # -------------------------------------------------
        # Make sure replacement exists
        # -------------------------------------------------

        if (
            replacement.lower()
            not in simplified_lower
        ):
            continue

        cleaned.append(
            {
                "word": word,
                "replacement": replacement
            }
        )

    # -----------------------------------------------------
    # Remove duplicates
    # -----------------------------------------------------

    unique = []

    seen = set()

    for item in cleaned:

        key = (
            item["word"].lower(),
            item["replacement"].lower()
        )

        if key in seen:
            continue

        seen.add(key)

        unique.append(
            item
        )

    return unique


# =========================================================
# EXTRACT UPLOADED FILE
# =========================================================

def extract_uploaded_file(uploaded_file):
    """
    Extract text from PDF, DOCX, or TXT uploads.

    Used directly by /api/simplify/ when the frontend sends
    a file in the same multipart/form-data request.
    """

    if not uploaded_file:
        return ""

    filename = uploaded_file.name.lower().strip()

    # -----------------------------------------------------
    # TXT
    # -----------------------------------------------------

    if filename.endswith(".txt"):

        try:
            return uploaded_file.read().decode("utf-8")

        except UnicodeDecodeError:

            uploaded_file.seek(0)

            return uploaded_file.read().decode("latin-1")

    # -----------------------------------------------------
    # DOCX
    # -----------------------------------------------------

    if filename.endswith(".docx"):

        document = Document(uploaded_file)

        paragraphs = []

        for paragraph in document.paragraphs:

            paragraph_text = paragraph.text.strip()

            if paragraph_text:
                paragraphs.append(paragraph_text)

        return "\n".join(paragraphs)

    # -----------------------------------------------------
    # PDF
    # -----------------------------------------------------

    if filename.endswith(".pdf"):

        reader = PdfReader(uploaded_file)

        pages = []

        for page in reader.pages:

            try:
                page_text = page.extract_text() or ""

            except Exception as page_error:

                print(
                    "PDF page extraction error:",
                    page_error
                )

                page_text = ""

            if page_text.strip():
                pages.append(page_text)

        return "\n".join(pages)

    raise ValueError(
        "Unsupported file type. "
        "Please upload a PDF, DOCX, or TXT file."
    )


# =========================================================
# SIMPLIFY TEXT
# =========================================================

@csrf_exempt
def simplify_text(request):

    # -----------------------------------------------------
    # POST only
    # -----------------------------------------------------

    if request.method != "POST":

        return JsonResponse(
            {
                "error":
                    "Only POST requests are allowed."
            },
            status=405
        )

    try:

        # =================================================
        # READ REQUEST
        # =================================================

        content_type = (
            request.content_type or ""
        )

        # -------------------------------------------------
        # JSON REQUEST
        # -------------------------------------------------

        if (
            "application/json"
            in content_type
        ):

            try:

                body = (
                    request.body
                    .decode("utf-8")
                )

                if not body.strip():

                    return JsonResponse(
                        {
                            "error":
                                "Request body is empty."
                        },
                        status=400
                    )

                data = json.loads(
                    body
                )

            except (
                json.JSONDecodeError,
                UnicodeDecodeError
            ):

                return JsonResponse(
                    {
                        "error":
                            "Invalid JSON request."
                    },
                    status=400
                )

            text = str(
                data.get(
                    "text",
                    ""
                )
            ).strip()

            level = str(
                data.get(
                    "level",
                    "Beginner"
                )
            ).strip()

        # -------------------------------------------------
        # FORM DATA
        # -------------------------------------------------

        else:

            text = request.POST.get(
                "text",
                ""
            ).strip()

            level = request.POST.get(
                "level",
                "Beginner"
            ).strip()

        # =================================================
        # NORMALIZE LEVEL
        # =================================================

        level_map = {
            "beginner": "Beginner",
            "intermediate": "Intermediate",
            "advanced": "Advanced"
        }

        level = level_map.get(
            level.lower(),
            "Beginner"
        )

        # =================================================
        # EXTRACT UPLOADED FILE WHEN TEXT IS EMPTY
        # =================================================
        #
        # The React frontend sends both "text" and "file".
        # For PDF/DOCX uploads, the text field can be empty,
        # so the backend must extract the document here.
        # =================================================

        if request.FILES.get("file"):

            uploaded_file = request.FILES.get("file")

            text = extract_uploaded_file(
                uploaded_file
            ).strip()

            text = re.sub(
                r"\n{3,}",
                "\n\n",
                text
            )

        # =================================================
        # VALIDATE TEXT
        # =================================================

        if not text:

            return JsonResponse(
                {
                    "error":
                        "Please enter some text to simplify."
                },
                status=400
            )

        # -------------------------------------------------
        # Maximum input size
        # -------------------------------------------------

        if len(text) > 50000:

            return JsonResponse(
                {
                    "error":
                        "The text is too long. "
                        "Please use a shorter document."
                },
                status=400
            )

        # =================================================
        # DEBUG
        # =================================================

        print(
            "\n"
            + "=" * 60
        )

        print(
            "READ EASE SIMPLIFICATION REQUEST"
        )

        print(
            "Reading level:",
            level
        )

        print(
            "Characters:",
            len(text)
        )

        print(
            "Words:",
            len(text.split())
        )

        print(
            "=" * 60
        )

        # =================================================
        # BEFORE READABILITY
        # =================================================

        before_readability = (
            calculate_readability(
                text
            )
        )

        # =================================================
        # GRAMMAR SCORE
        # =================================================

        try:

            grammar_result = grammar_score(
                text
            )

            grammar_result = float(
                grammar_result
            )

            grammar_result = max(
                min(
                    grammar_result,
                    100
                ),
                0
            )

        except Exception as grammar_error:

            print(
                "Grammar score error:",
                grammar_error
            )

            grammar_result = 0

        # =================================================
        # AI SIMPLIFICATION
        # =================================================

        print(
            "Sending text to simplification pipeline..."
        )

        simplified_result = (
            simplify_pipeline(
                text,
                level
            )
        )

        print(
            "Pipeline returned:"
        )

        print(
            simplified_result
        )

        # =================================================
        # EXTRACT PIPELINE RESULT
        # =================================================

        simplified = ""

        difficult_words = []

        changes = []

        # -------------------------------------------------
        # DICTIONARY RESPONSE
        # -------------------------------------------------

        if isinstance(
            simplified_result,
            dict
        ):

            simplified = (
                simplified_result.get(
                    "simplified_text",
                    ""
                )
            )

            difficult_words = (
                simplified_result.get(
                    "difficult_words",
                    []
                )
            )

            changes = (
                simplified_result.get(
                    "changes",
                    []
                )
            )

            # -------------------------------------------------
            # BACKWARD COMPATIBILITY
            #
            # If old pipeline does not have "changes",
            # difficult_words may contain mapping dictionaries.
            # Extract those mappings.
            # -------------------------------------------------

            if not changes:

                possible_changes = []

                if isinstance(
                    difficult_words,
                    list
                ):

                    for item in difficult_words:

                        if not isinstance(
                            item,
                            dict
                        ):
                            continue

                        word = item.get(
                            "word",
                            ""
                        )

                        replacement = item.get(
                            "replacement",
                            ""
                        )

                        if (
                            word
                            and replacement
                        ):

                            possible_changes.append(
                                {
                                    "word":
                                        word,

                                    "replacement":
                                        replacement
                                }
                            )

                changes = possible_changes

        # -------------------------------------------------
        # STRING RESPONSE
        # -------------------------------------------------

        else:

            simplified = (
                simplified_result
            )

        # =================================================
        # VALIDATE SIMPLIFIED TEXT
        # =================================================

        if simplified is None:

            simplified = ""

        simplified = str(
            simplified
        ).strip()

        if not simplified:

            return JsonResponse(
                {
                    "error":
                        "The AI returned empty text."
                },
                status=500
            )

        # =================================================
        # CLEAN AI TEXT
        # =================================================

        # Remove accidental Markdown code fences.

        simplified = re.sub(
            r"^```(?:text|json)?\s*",
            "",
            simplified,
            flags=re.IGNORECASE
        )

        simplified = re.sub(
            r"\s*```$",
            "",
            simplified
        )

        simplified = simplified.strip()

        # =================================================
        # CLEAN DIFFICULT WORDS
        # =================================================

        difficult_words = normalize_difficult_words(
            difficult_words
        )

        # =================================================
        # CLEAN ACTUAL CHANGES
        # =================================================

        changes = clean_word_mappings(
            changes,
            text,
            simplified
        )

        # =================================================
        # KEEP ONLY REAL REPLACEMENTS
        # =================================================
        #
        # A difficult word can remain unchanged. In that case
        # it belongs in the vocabulary list but NOT in changes.
        # =================================================

        changes_lookup = {
            (
                item["word"].lower(),
                item["replacement"].lower()
            )
            for item in changes
        }

        for item in difficult_words:

            word = item.get("word", "")
            replacement = item.get("replacement", "")

            if not replacement:
                continue

            if (
                word.lower(),
                replacement.lower()
            ) not in changes_lookup:

                item["replacement"] = ""

        # Add the validated replacement to the matching
        # vocabulary item.
        for change in changes:

            change_word = change["word"].lower()

            for item in difficult_words:

                if item["word"].lower() == change_word:

                    item["replacement"] = change["replacement"]

                    break

        # =================================================
        # AFTER READABILITY
        # =================================================

        after_readability = (
            calculate_readability(
                simplified
            )
        )

        # =================================================
        # READABILITY IMPROVEMENT
        # =================================================

        readability_improvement = round(
            after_readability[
                "reading_ease"
            ]
            -
            before_readability[
                "reading_ease"
            ],
            2
        )

        # =================================================
        # GRADE CHANGE
        # =================================================

        grade_change = round(
            before_readability[
                "grade_level"
            ]
            -
            after_readability[
                "grade_level"
            ],
            1
        )

        # =================================================
        # WORD COUNTS
        # =================================================

        original_word_count = len(
            re.findall(
                r"\b[a-zA-Z]+(?:'[a-zA-Z]+)?\b",
                text
            )
        )

        simplified_word_count = len(
            re.findall(
                r"\b[a-zA-Z]+(?:'[a-zA-Z]+)?\b",
                simplified
            )
        )

        # =================================================
        # READING TIME
        # =================================================

        # Approximate reading speed:
        # 200 words per minute.

        original_reading_time = max(
            round(
                original_word_count / 200
            ),
            1
        )

        simplified_reading_time = max(
            round(
                simplified_word_count / 200
            ),
            1
        )

        # =================================================
        # SENTENCE COUNTS
        # =================================================

        original_sentence_count = len(
            [
                sentence
                for sentence in re.split(
                    r"[.!?]+",
                    text
                )
                if sentence.strip()
            ]
        )

        simplified_sentence_count = len(
            [
                sentence
                for sentence in re.split(
                    r"[.!?]+",
                    simplified
                )
                if sentence.strip()
            ]
        )

        # =================================================
        # RESPONSE DATA
        # =================================================

        response_data = {

            # -------------------------------------------------
            # TEXT
            # -------------------------------------------------

            "original_text":
                text,

            "simplified_text":
                simplified,

            "reading_level":
                level,

            # -------------------------------------------------
            # READABILITY
            # -------------------------------------------------

            "before_readability":
                before_readability[
                    "reading_ease"
                ],

            "after_readability":
                after_readability[
                    "reading_ease"
                ],

            "before_grade":
                before_readability[
                    "grade_level"
                ],

            "after_grade":
                after_readability[
                    "grade_level"
                ],

            "readability_improvement":
                readability_improvement,

            "grade_change":
                grade_change,

            # -------------------------------------------------
            # GRAMMAR
            # -------------------------------------------------

            "grammar_score":
                f"{round(grammar_result)}%",

            # -------------------------------------------------
            # READING TIME
            # -------------------------------------------------

            "reading_time":
                f"{original_reading_time} min",

            "original_reading_time":
                f"{original_reading_time} min",

            "simplified_reading_time":
                f"{simplified_reading_time} min",

            # -------------------------------------------------
            # WORD / SENTENCE STATISTICS
            # -------------------------------------------------

            "original_word_count":
                original_word_count,

            "simplified_word_count":
                simplified_word_count,

            "original_sentence_count":
                original_sentence_count,

            "simplified_sentence_count":
                simplified_sentence_count,

            # -------------------------------------------------
            # DIFFICULT WORDS
            #
            # Used by frontend for yellow highlighting.
            # -------------------------------------------------

            "difficult_words":
                difficult_words,

            # -------------------------------------------------
            # ACTUAL CHANGES
            #
            # Used by frontend for cyan highlighting.
            # -------------------------------------------------

            "changes":
                changes
        }

        # =================================================
        # DEBUG RESPONSE
        # =================================================

        print(
            "\nFINAL SIMPLIFICATION RESPONSE:"
        )

        print(
            json.dumps(
                response_data,
                indent=2,
                ensure_ascii=False
            )
        )

        print(
            "=" * 60
        )

        # =================================================
        # RETURN
        # =================================================

        return JsonResponse(
            response_data,
            status=200
        )

    except Exception as error:

        print(
            "\nSIMPLIFICATION ERROR:"
        )

        print(
            repr(error)
        )

        return JsonResponse(
            {
                "error":
                    "Unable to simplify the text.",

                "details":
                    str(error)
            },
            status=500
        )


# =========================================================
# UPLOAD FILE
# =========================================================

@csrf_exempt
def upload_file(request):

    if request.method != "POST":

        return JsonResponse(
            {
                "error":
                    "Only POST requests are allowed."
            },
            status=405
        )

    try:

        uploaded_file = (
            request.FILES.get(
                "file"
            )
        )

        if not uploaded_file:

            return JsonResponse(
                {
                    "error":
                        "No file was uploaded."
                },
                status=400
            )

        filename = (
            uploaded_file.name
            .lower()
            .strip()
        )

        print(
            "Uploaded file:",
            filename
        )

        # =================================================
        # TXT
        # =================================================

        if filename.endswith(
            ".txt"
        ):

            try:

                text = (
                    uploaded_file
                    .read()
                    .decode("utf-8")
                )

            except UnicodeDecodeError:

                uploaded_file.seek(0)

                text = (
                    uploaded_file
                    .read()
                    .decode("latin-1")
                )

        # =================================================
        # DOCX
        # =================================================

        elif filename.endswith(
            ".docx"
        ):

            document = Document(
                uploaded_file
            )

            paragraphs = []

            for paragraph in (
                document.paragraphs
            ):

                paragraph_text = (
                    paragraph.text.strip()
                )

                if paragraph_text:

                    paragraphs.append(
                        paragraph_text
                    )

            text = "\n".join(
                paragraphs
            )

        # =================================================
        # PDF
        # =================================================

        elif filename.endswith(
            ".pdf"
        ):

            reader = PdfReader(
                uploaded_file
            )

            pages = []

            for page in reader.pages:

                try:

                    page_text = (
                        page.extract_text()
                    )

                except Exception as page_error:

                    print(
                        "PDF page extraction error:",
                        page_error
                    )

                    page_text = ""

                if page_text:

                    pages.append(
                        page_text
                    )

            text = "\n".join(
                pages
            )

        # =================================================
        # UNSUPPORTED
        # =================================================

        else:

            return JsonResponse(
                {
                    "error":
                        "Unsupported file type. "
                        "Please upload a PDF, DOCX, or TXT file."
                },
                status=400
            )

        # =================================================
        # CLEAN TEXT
        # =================================================

        text = text.strip()

        if not text:

            return JsonResponse(
                {
                    "error":
                        "Unable to extract text from the uploaded file."
                },
                status=400
            )

        # Normalize excessive blank lines.

        text = re.sub(
            r"\n{3,}",
            "\n\n",
            text
        )

        print(
            "Extracted characters:",
            len(text)
        )

        return JsonResponse(
            {
                "filename":
                    uploaded_file.name,

                "text":
                    text
            },
            status=200
        )

    except Exception as error:

        print(
            "File upload error:",
            repr(error)
        )

        return JsonResponse(
            {
                "error":
                    "Unable to read the uploaded file.",

                "details":
                    str(error)
            },
            status=500
        )


# =========================================================
# DOWNLOAD SIMPLIFIED TEXT
# =========================================================

@csrf_exempt
def download_simplified_text(request):

    if request.method != "POST":

        return JsonResponse(
            {
                "error":
                    "Only POST requests are allowed."
            },
            status=405
        )

    try:

        content_type = (
            request.content_type or ""
        )

        # =================================================
        # JSON
        # =================================================

        if (
            "application/json"
            in content_type
        ):

            try:

                body = (
                    request.body
                    .decode("utf-8")
                )

                if not body.strip():

                    return JsonResponse(
                        {
                            "error":
                                "Request body is empty."
                        },
                        status=400
                    )

                data = json.loads(
                    body
                )

            except (
                json.JSONDecodeError,
                UnicodeDecodeError
            ):

                return JsonResponse(
                    {
                        "error":
                            "Invalid JSON request."
                    },
                    status=400
                )

            text = str(
                data.get(
                    "text",
                    ""
                )
            ).strip()

            file_format = str(
                data.get(
                    "format",
                    "txt"
                )
            ).lower().strip()

        # =================================================
        # FORM DATA
        # =================================================

        else:

            text = request.POST.get(
                "text",
                ""
            ).strip()

            file_format = request.POST.get(
                "format",
                "txt"
            ).lower().strip()

        # =================================================
        # VALIDATE
        # =================================================

        if not text:

            return JsonResponse(
                {
                    "error":
                        "No simplified text provided."
                },
                status=400
            )

        # =================================================
        # TXT
        # =================================================

        if file_format == "txt":

            response = HttpResponse(
                text,
                content_type=(
                    "text/plain; charset=utf-8"
                )
            )

            response[
                "Content-Disposition"
            ] = (
                'attachment; '
                'filename="simplified_text.txt"'
            )

            return response

        # =================================================
        # DOCX
        # =================================================

        if file_format == "docx":

            document = Document()

            document.add_heading(
                "ReadEase AI - Simplified Text",
                level=1
            )

            paragraphs = text.split(
                "\n"
            )

            for paragraph in paragraphs:

                paragraph = (
                    paragraph.strip()
                )

                if paragraph:

                    document.add_paragraph(
                        paragraph
                    )

            file_stream = io.BytesIO()

            document.save(
                file_stream
            )

            file_stream.seek(0)

            response = HttpResponse(
                file_stream.getvalue(),
                content_type=(
                    "application/vnd.openxmlformats-officedocument."
                    "wordprocessingml.document"
                )
            )

            response[
                "Content-Disposition"
            ] = (
                'attachment; '
                'filename="simplified_text.docx"'
            )

            return response

        # =================================================
        # PDF
        # =================================================

        if file_format == "pdf":

            file_stream = io.BytesIO()

            pdf = canvas.Canvas(
                file_stream,
                pagesize=A4
            )

            width, height = A4

            x = 50
            y = height - 50

            # -------------------------------------------------
            # Title
            # -------------------------------------------------

            pdf.setFont(
                "Helvetica-Bold",
                16
            )

            pdf.drawString(
                x,
                y,
                "ReadEase AI - Simplified Text"
            )

            y -= 35

            # -------------------------------------------------
            # Body
            # -------------------------------------------------

            pdf.setFont(
                "Helvetica",
                11
            )

            max_width = (
                width - 100
            )

            paragraphs = text.split(
                "\n"
            )

            for paragraph in paragraphs:

                words = paragraph.split()

                line = ""

                for word in words:

                    test_line = (
                        f"{line} {word}"
                    ).strip()

                    line_width = (
                        pdf.stringWidth(
                            test_line,
                            "Helvetica",
                            11
                        )
                    )

                    if line_width <= max_width:

                        line = test_line

                    else:

                        if line:

                            pdf.drawString(
                                x,
                                y,
                                line
                            )

                            y -= 18

                        line = word

                    # -----------------------------------------
                    # Page break
                    # -----------------------------------------

                    if y < 50:

                        pdf.showPage()

                        pdf.setFont(
                            "Helvetica",
                            11
                        )

                        y = height - 50

                if line:

                    pdf.drawString(
                        x,
                        y,
                        line
                    )

                    y -= 18

                y -= 8

                if y < 50:

                    pdf.showPage()

                    pdf.setFont(
                        "Helvetica",
                        11
                    )

                    y = height - 50

            pdf.save()

            file_stream.seek(0)

            response = HttpResponse(
                file_stream.getvalue(),
                content_type="application/pdf"
            )

            response[
                "Content-Disposition"
            ] = (
                'attachment; '
                'filename="simplified_text.pdf"'
            )

            return response

        # =================================================
        # INVALID FORMAT
        # =================================================

        return JsonResponse(
            {
                "error":
                    "Unsupported download format."
            },
            status=400
        )

    except Exception as error:

        print(
            "Download error:",
            repr(error)
        )

        return JsonResponse(
            {
                "error":
                    "Unable to create the file.",

                "details":
                    str(error)
            },
            status=500
        )


# =========================================================
# SEND VERIFICATION OTP
# =========================================================

@api_view(["POST"])
def send_verification_otp(request):

    email = str(
        request.data.get(
            "email",
            ""
        )
    ).strip().lower()

    # =================================================
    # VALIDATE EMAIL
    # =================================================

    email_pattern = (
        r"^[a-zA-Z0-9._%+-]+@"
        r"[a-zA-Z0-9-]+"
        r"(\.[a-zA-Z0-9-]+)+$"
    )

    if not re.match(
        email_pattern,
        email
    ):

        return Response(
            {
                "success": False,

                "message":
                    "Please enter a valid email address."
            },
            status=status.HTTP_400_BAD_REQUEST
        )

    # =================================================
    # GENERATE OTP
    # =================================================

    otp = generate_otp()

    # =================================================
    # STORE OTP
    # =================================================

    store_otp(
        email,
        otp
    )

    # =================================================
    # SEND EMAIL
    # =================================================

    try:

        send_mail(
            subject=(
                "ReadEase AI - Email Verification Code"
            ),

            message=(
                "Hello,\n\n"
                "Your ReadEase AI verification code is:\n\n"
                f"{otp}\n\n"
                "This code will expire in 10 minutes.\n\n"
                "If you did not request this code, "
                "you can safely ignore this email.\n\n"
                "Regards,\n"
                "ReadEase AI Team"
            ),

            from_email=None,

            recipient_list=[
                email
            ],

            fail_silently=False
        )

        return Response(
            {
                "success": True,

                "message":
                    "Verification code sent successfully."
            },
            status=status.HTTP_200_OK
        )

    except Exception as error:

        print(
            "Email sending error:",
            repr(error)
        )

        return Response(
            {
                "success": False,

                "message":
                    "Unable to send verification email."
            },
            status=status.HTTP_500_INTERNAL_SERVER_ERROR
        )


# =========================================================
# VERIFY EMAIL OTP
# =========================================================

@api_view(["POST"])
def verify_verification_otp(request):

    email = str(
        request.data.get(
            "email",
            ""
        )
    ).strip().lower()

    entered_otp = str(
        request.data.get(
            "otp",
            ""
        )
    ).strip()

    # =================================================
    # VALIDATE INPUT
    # =================================================

    if (
        not email
        or not entered_otp
    ):

        return Response(
            {
                "success": False,

                "message":
                    "Email and verification code are required."
            },
            status=status.HTTP_400_BAD_REQUEST
        )

    # =================================================
    # VERIFY OTP
    # =================================================

    try:

        verified, message = (
            verify_otp(
                email,
                entered_otp
            )
        )

    except Exception as error:

        print(
            "OTP verification error:",
            repr(error)
        )

        return Response(
            {
                "success": False,

                "message":
                    "Unable to verify the code right now."
            },
            status=status.HTTP_500_INTERNAL_SERVER_ERROR
        )

    # =================================================
    # SUCCESS
    # =================================================

    if verified:

        return Response(
            {
                "success": True,

                "message":
                    message
            },
            status=status.HTTP_200_OK
        )

    # =================================================
    # FAILURE
    # =================================================

    return Response(
        {
            "success": False,

            "message":
                message
        },
        status=status.HTTP_400_BAD_REQUEST
    )